import io
from django.http import FileResponse
from docx import Document
from content.models import Plant
from docxcompose.composer import Composer

def replace_full_text(doc, placeholder, source_doc):
    for para in doc.paragraphs:
        if placeholder in para.text:
            para.text = "" 
            for source_para in source_doc.paragraphs:
                new_para = para.insert_paragraph_before()
                for run in source_para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.size = run.font.size
                    if run.font.color and run.font.color.rgb:
                        new_run.font.color.rgb = run.font.color.rgb
                    if run.font.name:
                        new_run.font.name = run.font.name

            return
        
def replace_text_with_format(paragraph, placeholder, replacement):
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)

def generate_docx(plant: Plant, template_path="static/docs/template1.docx", template_end_path="static/docs/template2.docx"):
    buffer = io.BytesIO()

    doc = Document(template_path)
    genus_description = plant.genus.genus_description[0].lower() + plant.genus.genus_description[1:]
    replacements = {
        "plant_name": f"{plant.plant_name}",
        "plant_description": f"{plant.plant_description}",
        "genus_name": f"{plant.genus.genus_name}",
        "genus_details": f"{genus_description}",
        "supplier_name": f"{plant.supplier.supplier_name}",
        "contact_details": f"{plant.supplier.contact_details}",
    }
    for para in doc.paragraphs:
        for placeholder, replacement in replacements.items():
            replace_text_with_format(para, placeholder, replacement)

    composer = Composer(doc)
    plant_details_doc = Document(f"static/docs/{plant.additional_info}")
    composer.append(plant_details_doc)

    end_doc = Document(template_end_path)
    for para in end_doc.paragraphs:
        for placeholder, replacement in replacements.items():
            replace_text_with_format(para, placeholder, replacement)
    composer.append(end_doc)

    doc.save(buffer)
    buffer.seek(0)
    filename = f"{plant.plant_name}.docx"
    response = FileResponse(buffer, as_attachment=True, filename=filename)
    return response
